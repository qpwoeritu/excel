"""Generator dokumentácie k výpočtovému programu (VBA / Excel)."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Calibri'
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.autofit = False
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        set_cell_shading(hdr_cells[i], '1F4E79')
        for r in hdr_cells[i].paragraphs[0].runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ''
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def add_placeholder(doc, label):
    """Pridáva miesto na vloženie obrázka (rámček)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'[ Sem vložte screenshot: {label} ]')
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


doc = Document()

# default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# nastavenie okrajov stránky
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)

# =======================================================
# TITULNÁ STRANA
# =======================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('\n\n\n\nTechnická dokumentácia\nvýpočtového programu')
run.bold = True
run.font.size = Pt(28)
run.font.name = 'Calibri'

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('\nVýpočet ustáleného chodu a skratových pomerov\n'
                  'v elektrizačných sieťach (Excel / VBA)')
run.font.size = Pt(16)
run.italic = True

doc.add_paragraph('\n\n\n')

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Dokument popisuje matematický model, vstupné a výstupné\n'
                   'údaje, modelovanie jednotlivých prvkov siete a algoritmus\n'
                   'výpočtu implementovaný v VBA module Microsoft Excel.')
run.font.size = Pt(11)

doc.add_page_break()

# =======================================================
# 1. ÚVOD
# =======================================================
add_heading(doc, '1. Úvod a účel programu', level=1)

add_para(doc,
    'Výpočtový program je implementovaný v jazyku VBA (Visual Basic for '
    'Applications) ako rozšírenie zošita Microsoft Excel. Slúži na riešenie '
    'dvoch základných úloh elektroenergetiky:')
add_bullet(doc, 'výpočet ustáleného chodu (load-flow) iteračnou metódou Newton–Raphson;')
add_bullet(doc, 'výpočet trojfázového skratu Ik3″ vo všetkých uzloch siete '
                'metódou inverzie skratovej admitančnej matice.')

add_para(doc,
    'Vstupné údaje aj vypočítané výsledky sú uložené priamo v hárkoch zošita; '
    'výsledky sú navyše doplnené do jednolinkovej schémy (SLD) prostredníctvom '
    'mechanizmu pomenovaných „tagov". Program podporuje siete s rôznymi '
    'napäťovými hladinami (VVN/VN/NN), modelovanie izolovaných prvkov a '
    'detekciu nepripojených častí siete topologickou analýzou (BFS).')

add_heading(doc, '1.1 Štruktúra zdrojového kódu', level=2)
add_table(doc,
    ['Modul', 'Účel'],
    [
        ('modTypes.bas',      'Definícia dátových typov (Complex, BusType) a konštánt.'),
        ('modComplex.bas',    'Knižnica komplexnej aritmetiky (sčítanie, násobenie, atan2…).'),
        ('modIO.bas',         'Načítavanie vstupov z hárkov a zápis výsledkov.'),
        ('modUtils.bas',      'Pomocné funkcie (parsing čísel, vyhľadávanie uzlov, prácu s hárkami).'),
        ('modTopology.bas',   'Topologická analýza – BFS detekcia izolovaných prvkov.'),
        ('modYBus.bas',       'Zostavenie admitančnej matice Y pre load-flow.'),
        ('modNR.bas',         'Newton–Raphsonov iteračný riešič tokov výkonov.'),
        ('modShortCircuit.bas','Zostavenie skratovej matice a výpočet Ik3″.'),
        ('modSLD.bas',        'Zápis výsledkov do jednolinkovej schémy podľa tagov.'),
        ('modProgress.bas',   'Stavový panel, časovač a obsluha klávesy ESC.'),
        ('modMain.bas',       'Vstupný bod runCALC – orchestrácia fáz výpočtu.'),
    ],
    col_widths=[4.5, 11.5])

doc.add_page_break()

# =======================================================
# 2. VŠEOBECNÉ VLASTNOSTI VÝPOČTU
# =======================================================
add_heading(doc, '2. Všeobecné vlastnosti výpočtu', level=1)

add_heading(doc, '2.1 Pomerné jednotky (per-unit)', level=2)
add_para(doc,
    'Program pracuje interne výhradne v pomerných jednotkách. Vzťažné hodnoty:')
add_bullet(doc, 'Bázový výkon Sbase [MVA] – spoločný pre celú sieť, číta sa z hárka index.')
add_bullet(doc, 'Bázové napätie Ubase [kV] – individuálne pre každý uzol podľa '
                'jeho menovitej napäťovej hladiny (zoznam hladín VLevels z hárka data).')
add_bullet(doc, 'Bázová impedancia Zbase = Ubase² / Sbase [Ω].')
add_bullet(doc, 'Bázový prúd Ibase = Sbase / (√3 · Ubase) [kA].')
add_para(doc,
    'Všetky vstupné impedancie (Ω), admitancie (S) a výkony (MW, MVAr) sa pri '
    'načítaní automaticky prepočítavajú do p.u. Výstupné hodnoty sú prepočítané '
    'späť do fyzikálnych jednotiek (kV, A, MW, MVAr, kW strát).')

add_heading(doc, '2.2 Typy uzlov', level=2)
add_table(doc,
    ['Kód', 'Typ', 'Známe veličiny', 'Použitie'],
    [
        ('0', 'Slack (referenčný)', '|V|, θ = 0', 'Napájací bod siete, vyrovnáva bilanciu.'),
        ('1', 'PQ',                  'P, Q',       'Spotrebiteľské uzly, výmena.'),
        ('2', 'PV',                  'P, |V|',     'Generátorové uzly s reguláciou napätia.'),
    ],
    col_widths=[1.5, 4.0, 4.0, 6.5])

add_heading(doc, '2.3 Riadiace parametre (hárok index)', level=2)
add_table(doc,
    ['Bunka', 'Význam'],
    [
        ('G5', 'Režim výpočtu: 1 = load-flow, 2 = skraty.'),
        ('I3:J8', 'Stavový panel jednotlivých fáz (Načítanie, Tvorba matice, '
                  'NR, Skrat, SLD) – farba bunky indikuje stav (čaká / prebieha / hotovo / chyba).'),
        ('Sbase', 'Bázový výkon siete (MVA), prevzatý z hárka index/data.'),
        ('Epsilon', 'Kritérium konvergencie NR (default 1·10⁻⁶ p.u.).'),
        ('Max. iterácií', 'Bezpečnostná hranica iterácií NR (default 20).'),
    ],
    col_widths=[3.0, 13.0])

add_placeholder(doc, 'hárok „index" s riadiacim panelom a tlačidlom spustenia')

doc.add_page_break()

# =======================================================
# 3. MODELOVANIE PRVKOV
# =======================================================
add_heading(doc, '3. Modelovanie jednotlivých prvkov siete', level=1)

# --- 3.1 Uzly
add_heading(doc, '3.1 Uzly (hárok „uzly")', level=2)
add_para(doc,
    'Každý riadok hárka uzly reprezentuje jeden uzol siete. Uzol je '
    'matematicky charakterizovaný komplexným napätím V = |V|·e^(jθ). '
    'Podľa typu (Slack / PQ / PV) sú niektoré jeho parametre známe a iné '
    'sú výsledkom výpočtu.')

add_para(doc, 'Štruktúra vstupných stĺpcov:', bold=True)
add_table(doc,
    ['Stĺpec', 'Veličina', 'Jednotka', 'Poznámka'],
    [
        ('B', 'Názov uzla',         '–',     'Jednoznačný identifikátor (použitý vo všetkých prvkoch).'),
        ('C', 'Typ',                '0/1/2', '0 = Slack, 1 = PQ, 2 = PV.'),
        ('D', 'Počiatočné |V|',     'kV',    'Pre Slack a PV pevné, pre PQ počiatočný odhad.'),
        ('F', 'Činný výkon P',      'MW',    'Spotreba kladná; výroba záporná (alebo opačná konvencia podľa zadania).'),
        ('G', 'Jalový výkon Q',     'Mvar',  'Pre PQ pevné; pre PV počiatočný odhad.'),
    ],
    col_widths=[1.5, 4.0, 2.0, 8.5])

add_para(doc, 'Výstupy zapísané do hárka uzly:', bold=True)
add_table(doc,
    ['Stĺpec', 'Veličina', 'Jednotka'],
    [
        ('H', '|V| vypočítané', 'kV'),
        ('I', 'θ vypočítaný',   '°'),
    ],
    col_widths=[2.0, 6.0, 4.0])

add_placeholder(doc, 'hárok „uzly" so vstupmi a vypočítanými stĺpcami H, I')

# --- 3.2 Vedenia
add_heading(doc, '3.2 Vedenia (hárok „vedenia")', level=2)
add_para(doc,
    'Vedenia sú modelované π-článkom: pozdĺžnou impedanciou Z = R + jX a '
    'priečnou susceptanciou B/2 na oboch koncoch.')

p = doc.add_paragraph()
run = p.add_run('       From  ──[ R + jX ]──  To\n'
                '              │             │\n'
                '            jB/2          jB/2\n'
                '              │             │\n'
                '             ─┴─           ─┴─')
run.font.name = 'Consolas'
run.font.size = Pt(10)

add_para(doc, 'Vstupné stĺpce (relevantné):', bold=True)
add_table(doc,
    ['Stĺpec', 'Veličina', 'Jednotka', 'Poznámka'],
    [
        ('C',     'Názov vedenia',       '–',  ''),
        ('D, E',  'Uzol From, Uzol To',  '–',  'Názvy musia existovať v hárku „uzly".'),
        ('M',     'Sériové R',            'Ω',  ''),
        ('N',     'Sériové X',            'Ω',  ''),
        ('O',     'Priečna B (celkom)',  'S',  'Polovica sa dáva na každý koniec.'),
    ],
    col_widths=[1.5, 4.5, 2.0, 8.0])

add_para(doc, 'Vypočítané výstupy (stĺpce X–AB):', bold=True)
add_bullet(doc, '|I_ij| – modul prúdu vo vedení [A];')
add_bullet(doc, 'ΔU – relatívny úbytok napätia medzi koncami [%];')
add_bullet(doc, 'P_ij, Q_ij – tok činného a jalového výkonu na začiatku vedenia [MW, MVAr];')
add_bullet(doc, 'P_str – činné straty na vedení [kW].')

add_placeholder(doc, 'hárok „vedenia" – vstupné parametre a vypočítané stĺpce')

# --- 3.3 Transformátory
add_heading(doc, '3.3 Transformátory (hárok „transformatory")', level=2)
add_para(doc,
    'Trojfázové dvojvinuťové transformátory sa modelujú ekvivalentným π-modelom '
    's prevodom a (turn-ratio). Pozdĺžna vetva obsahuje sériovú impedanciu '
    'Z = R + jX, priečna vetva obsahuje magnetizačné straty G + jB '
    '(reprezentujúce straty v železe a magnetizačný prúd). Prevod a sa aplikuje '
    'ako ideálny transformátor zaradený do schémy.')

add_para(doc, 'Vstupné údaje:', bold=True)
add_table(doc,
    ['Stĺpec', 'Veličina', 'Jednotka', 'Poznámka'],
    [
        ('C',  'Názov',                  '–',  ''),
        ('D, E', 'Uzol primár / sekundár', '–', ''),
        ('R',  'Režim ZK',               'text', 'Ak obsahuje „zk" → predpokladá sa zapojenie zo strany ZK.'),
        ('R/S', 'R sériové',             'Ω',  'Pozícia závisí od režimu zapojenia.'),
        ('S/T', 'X sériové',             'Ω',  ''),
        ('T/U', 'G magnetizačné',        'S',  'Straty v železe.'),
        ('U/V', 'B magnetizačné',        'S',  ''),
        ('V/W', 'Prevod a',              '–',  'Bezrozmerný, typicky 0.95–1.05.'),
    ],
    col_widths=[1.5, 5.0, 2.0, 7.5])

add_para(doc, 'Vypočítané výstupy:', bold=True)
add_bullet(doc, '|I_prim|, |I_sec| – moduly prúdu na primárnej a sekundárnej strane [A];')
add_bullet(doc, 'P_prim, Q_prim a P_sec, Q_sec – toky výkonu na oboch stranách [MW, MVAr];')
add_bullet(doc, 'P_str – činné straty transformátora [kW] (vinutie + železo).')

add_placeholder(doc, 'hárok „transformatory" – vstupy a výstupy výpočtu')

# --- 3.4 Reaktory
add_heading(doc, '3.4 Reaktory a diferenciálne reaktory '
                 '(hárky „reaktory", „dif_reaktory")', level=2)
add_para(doc,
    'Reaktor je modelovaný čisto pozdĺžnou impedanciou Z = R + jX, bez priečnej '
    'kapacity. Diferenciálne (sekciovacie) reaktory majú rovnakú topologickú '
    'štruktúru, ale slúžia na obmedzenie skratových prúdov medzi sekciami '
    'rozvodne; v matici Y vystupujú rovnako.')

add_table(doc,
    ['Stĺpec', 'Veličina', 'Jednotka'],
    [
        ('C',    'Názov',          '–'),
        ('D, E', 'Uzol From, To',  '–'),
        ('H',    'R sériové',      'Ω'),
        ('I',    'X sériové',      'Ω'),
    ],
    col_widths=[1.5, 4.5, 3.0])

add_para(doc, 'Výstupy: |I|, ΔU, P, Q a Ploss obdobne ako pri vedeniach.')

add_placeholder(doc, 'hárky „reaktory" / „dif_reaktory"')

# --- 3.5 Kompenzácie
add_heading(doc, '3.5 Kompenzácie (hárok „kompenzácia")', level=2)
add_para(doc,
    'Kompenzačné prvky (kondenzátorové batérie, tlmivky) sú modelované ako '
    'priečna susceptancia pripojená medzi uzol a zem. Výpočet výslednej '
    'reaktancie:')
p = doc.add_paragraph()
run = p.add_run('   X_net = X_C − X_L      →      B = 1 / X_net   [S]')
run.font.name = 'Consolas'
run.font.size = Pt(10)
add_para(doc,
    'Ak X_L > 0 a X_C = 0 dostávame čisto induktívnu tlmivku (B < 0), '
    'ak X_L = 0 a X_C > 0 dostávame čisto kapacitnú batériu (B > 0). '
    'Stĺpec N obsahuje Status: 1 = aktívna, 0 = odpojená (B = 0).')

add_table(doc,
    ['Stĺpec', 'Veličina', 'Jednotka'],
    [
        ('B', 'Názov',     '–'),
        ('C', 'Uzol',      '–'),
        ('N', 'Status',    '0/1'),
        ('P', 'X_L',       'Ω'),
        ('Q', 'X_C',       'Ω'),
    ],
    col_widths=[1.5, 4.5, 3.0])

add_para(doc, 'Výstup: napätie U na kompenzácii [kV] (stĺpec O).')

add_placeholder(doc, 'hárok „kompenzácia"')

# --- 3.6 Motory VN
add_heading(doc, '3.6 Asynchrónne motory VN (hárok „motoryVN")', level=2)
add_para(doc,
    'V load-flow sa motor reprezentuje ako pevná priečna admitancia Y = G + jB '
    'pripojená do svorkového uzla – modeluje sa tým činný odber a jalová '
    'spotreba motora. Pri skratovej úlohe je motor nahradený subtransientnou '
    'reaktanciou Xk pripojenou medzi svorkový uzol a vnútorný uzol s nominálnym '
    'napätím (príspevok do skratového prúdu zo strany motora).')

add_table(doc,
    ['Stĺpec', 'Veličina', 'Jednotka', 'Použitie'],
    [
        ('B', 'Názov',       '–',   ''),
        ('C', 'Uzol',        '–',   ''),
        ('L', 'R',           'Ω',   'Odpor (informačný, súčasť Z motora).'),
        ('P', 'X_k subtran.', 'Ω',  'Použité pri skrate.'),
        ('Q', 'G',           'S',   'Použité v load-flow.'),
        ('R', 'B',           'S',   'Použité v load-flow.'),
        ('S', 'Status',      '0/1', '0 = vypnutý motor (vylučuje sa).'),
    ],
    col_widths=[1.5, 4.0, 2.0, 8.5])

add_para(doc, 'Výstup: |I| na svorkách motora [A] a činné straty [kW].')

add_placeholder(doc, 'hárok „motoryVN"')

# --- 3.7 Generátory
add_heading(doc, '3.7 Generátory (hárok „generatory")', level=2)
add_para(doc,
    'Program podporuje dva spôsoby modelovania generátora:')
add_bullet(doc, 'Režim PQ – generátor je pevná injekcia (P_gen, Q_ref) do svorkového uzla. '
                'Vhodný pre rozvodne s pevne stanoveným pracovným bodom.')
add_bullet(doc, 'Režim EMF – fyzikálny model synchrónneho stroja: vnútorné napätie |E| '
                'za reaktanciou Ra + jXs. Riešič automaticky doplní fantómový PV uzol s '
                'pevným |E|, ktorý je k svorkovému uzlu pripojený sériovou vetvou Ra + jXs. '
                'Hodnoty |E| a vnútorný činný výkon P_int sa kalibrujú z referenčného bodu '
                '(V_ref, Q_ref, P_gen).')

add_table(doc,
    ['Stĺpec', 'Veličina', 'Jednotka', 'Použitie'],
    [
        ('B', 'Názov',       '–',     ''),
        ('C', 'Svorka (uzol)','–',    'Uzol, kam je generátor pripojený.'),
        ('D', 'Režim',       'PQ/EMF',''),
        ('E', 'Status',      '0/1',   ''),
        ('L', 'Ra',          'Ω',     'Odpor armatúry (iba EMF).'),
        ('M', 'P_gen',       'MW',    'Činný výkon generátora.'),
        ('N', 'Q_ref',       'Mvar',  'Referenčný jalový výkon (PQ – pevný, EMF – kalibračný).'),
        ('O', 'Xs',          'Ω',     'Synchrónna reaktancia (iba EMF).'),
        ('P', 'X_d″',        'Ω',     'Subtransientná reaktancia (skraty).'),
        ('Q', 'V_ref',       'kV',    'Referenčné napätie svorky (EMF).'),
    ],
    col_widths=[1.5, 4.0, 2.0, 8.5])

add_para(doc, 'Vypočítané výstupy:', bold=True)
add_bullet(doc, 'δ – uhol vnútorného napätia (rotorový uhol) [°];')
add_bullet(doc, 'Q_gen – výsledný jalový výkon generátora [Mvar];')
add_bullet(doc, '|I| – modul prúdu na svorkách [A];')
add_bullet(doc, 'P_loss – činné straty na vinutí (Ra·I²) [kW].')

add_placeholder(doc, 'hárok „generatory"')

# --- 3.8 Spínače
add_heading(doc, '3.8 Spínače (hárok „spinace")', level=2)
add_para(doc,
    'Spínač je modelovaný ako vetva s malou impedanciou Z = R + jX medzi dvomi '
    'uzlami. Pri rozpojenom stave (Status = 0) je odpojený od matice Y, čo má '
    'priamy vplyv aj na topologickú analýzu – spínač sa nepoužije pri BFS '
    'a uzly za ním môžu byť označené ako izolované.')

add_table(doc,
    ['Stĺpec', 'Veličina', 'Jednotka'],
    [
        ('C',    'Názov',          '–'),
        ('D, E', 'Uzol From, To',  '–'),
        ('G',    'R',              'Ω'),
        ('H',    'X',              'Ω'),
        ('Status', 'Stav',         '0/1'),
    ],
    col_widths=[2.0, 4.5, 3.0])

add_para(doc, 'Výstup: |I| na spínači [A].')

add_placeholder(doc, 'hárok „spinace"')

doc.add_page_break()

# =======================================================
# 4. ALGORITMUS LOAD-FLOW
# =======================================================
add_heading(doc, '4. Algoritmus výpočtu ustáleného chodu', level=1)

add_heading(doc, '4.1 Admitančná matica Y', level=2)
add_para(doc,
    'Modul modYBus.bas zostaví komplexnú admitančnú maticu Y o rozmere n × n, '
    'kde n je celkový počet uzlov (vrátane fantómových PV uzlov EMF generátorov). '
    'Príspevky jednotlivých prvkov:')
add_bullet(doc, 'vedenie (i,j): y_ij = 1/(R+jX), priečna kapacita B/2 sa pridá '
                'na diagonálu Y[i,i] a Y[j,j];')
add_bullet(doc, 'transformátor: rovnaký π-model s prevodom – Y[i,i] += y/a², '
                'Y[j,j] += y, Y[i,j] -= y/a, Y[j,i] -= y/a;')
add_bullet(doc, 'reaktor / dif. reaktor / spínač: čistá sériová admitancia;')
add_bullet(doc, 'kompenzácia: príspevok jB na diagonálu uzla;')
add_bullet(doc, 'motor: admitancia G + jB na diagonálu uzla;')
add_bullet(doc, 'EMF generátor: sériová vetva Ra + jXs medzi fantómovým a svorkovým uzlom.')
add_para(doc,
    'Izolované uzly (zistené BFS) dostávajú jednotkovú diagonálnu admitanciu, '
    'aby bola matica regulárna a riešič ostal numericky stabilný.')

add_heading(doc, '4.2 Newton–Raphsonov iteračný cyklus', level=2)
add_para(doc, 'Pre každý uzol i sa počíta činný a jalový výkon vzťahmi:')
p = doc.add_paragraph()
run = p.add_run(
    '   P_i = |V_i| Σ_k |V_k| · ( G_ik · cos θ_ik + B_ik · sin θ_ik )\n'
    '   Q_i = |V_i| Σ_k |V_k| · ( G_ik · sin θ_ik − B_ik · cos θ_ik )')
run.font.name = 'Consolas'
run.font.size = Pt(10)
add_para(doc,
    'V každej iterácii sa zostaví vektor nesúladu ΔP, ΔQ a Jacobiho matica '
    'J = [∂P/∂θ, ∂P/∂|V| ; ∂Q/∂θ, ∂Q/∂|V|]. Sústava J·Δx = [ΔP; ΔQ] sa rieši '
    'Gaussovou elimináciou s čiastočným pivotovaním. Stav sa aktualizuje '
    'V ← V + ΔV, θ ← θ + Δθ. Cyklus končí ak max(|ΔP|, |ΔQ|) < ε, kde ε je '
    'zadané v hárku index (typicky 10⁻⁶ p.u.).')

add_para(doc, 'Vlastnosti riešiča:', bold=True)
add_bullet(doc, 'analytická Jacobiho matica (žiadna numerická diferenciácia);')
add_bullet(doc, 'kvadratická konvergencia v okolí riešenia;')
add_bullet(doc, 'maximálny počet iterácií ako poistka proti divergencii;')
add_bullet(doc, 'priebežné logovanie napätí a hodnoty ε do hárkov „napatia" a „epsilon";')
add_bullet(doc, 'možnosť prerušenia výpočtu stlačením ESC (kontrolované cez DoEvents).')

add_placeholder(doc, 'hárky „napatia" a „epsilon" – iteračný priebeh')

# =======================================================
# 5. SKRATOVÝ VÝPOČET
# =======================================================
add_heading(doc, '5. Skratový výpočet (Ik3″)', level=1)

add_para(doc,
    'Modul modShortCircuit.bas zostavuje samostatnú skratovú admitančnú maticu '
    'Y_sc, v ktorej sú všetky pasívne prvky reprezentované svojimi '
    'subtransientnými parametrami:')
add_bullet(doc, 'vedenia, transformátory a reaktory – ako pri load-flow (R + jX, B);')
add_bullet(doc, 'motory – ako čisto reaktívna admitancia y = 1/(jX_k);')
add_bullet(doc, 'generátory – ako vetva Ra + jX_d″ od fantómového uzla so zdrojom EMF;')
add_bullet(doc, 'slack (napájacia sieť) – ekvivalentná zdrojová impedancia odvodená '
                'zo zadaného Ik3 siete vzorcom Z_grid = 1.1 · Sbase / (√3 · U · Ik3).')

add_para(doc,
    'Inverziou Y_sc sa získa matica Theveninových impedancií Z_th. Skratový '
    'prúd v uzle i sa potom vypočíta:')
p = doc.add_paragraph()
run = p.add_run('   I_k″ = (c · U_n) / (√3 · |Z_th,i|),     c = 1.1')
run.font.name = 'Consolas'
run.font.size = Pt(10)
add_para(doc,
    'Koeficient c = 1.1 zodpovedá norme STN EN 60909 pre maximálne počiatočné '
    'skratové prúdy. Výsledky sa zapisujú do hárka uzly (stĺpec Ik3″ [kA]).')

add_para(doc,
    'Inverzia komplexnej matice je riešená Gauss–Jordanovou elimináciou s '
    'čiastočným pivotovaním priamo nad UDT typom Complex (bez rozkladu na '
    'reálnu 2n × 2n maticu), čo zaisťuje rádovo lepší výkon aj presnosť.')

add_placeholder(doc, 'výsledky skratu v hárku „uzly" (stĺpec Ik3″)')

doc.add_page_break()

# =======================================================
# 6. TOPOLOGICKÁ ANALÝZA
# =======================================================
add_heading(doc, '6. Topologická analýza siete', level=1)
add_para(doc,
    'Pred zostavením matice Y modul modTopology.bas vykoná BFS prehľadávanie '
    'zo Slack uzla. Hrany grafu tvoria všetky pripojené impedančné prvky '
    '(vedenia, trafá, reaktory) a zopnuté spínače. Uzly, ktoré nie sú BFS '
    'dosiahnuteľné, sa označia ako izolované a:')
add_bullet(doc, 'v hárku uzly dostanú v stĺpcoch H, I hodnoty „izolovane" / „-";')
add_bullet(doc, 'sú vyradené z výpočtu (do Y matice dostávajú jednotku na diagonálu);')
add_bullet(doc, 'sú zapísané do hárka „report" v zoznamoch „Izolované uzly", '
                '„Izolované vedenia", „Izolované trafá", „Izolované kompenzácie".')

add_placeholder(doc, 'hárok „report" so zoznamom izolovaných prvkov')

# =======================================================
# 7. VÝSTUPY
# =======================================================
add_heading(doc, '7. Prehľad výstupov programu', level=1)

add_heading(doc, '7.1 Doplnené stĺpce vo vstupných hárkoch', level=2)
add_table(doc,
    ['Hárok', 'Doplnené stĺpce', 'Veličiny'],
    [
        ('uzly',          'H, I (+ Ik3″)', '|V| [kV], θ [°], Ik3″ [kA]'),
        ('vedenia',       'Q–U',           '|I| [A], ΔU [%], P [MW], Q [MVAr], P_str [kW]'),
        ('transformatory','X–AD',          'I_prim, I_sec, P_prim, Q_prim, P_sec, Q_sec, P_str'),
        ('reaktory',      'Z–AD',          '|I|, ΔU, P, Q, P_str'),
        ('dif_reaktory',  'X–AB',          '|I|, ΔU, P, Q, P_str'),
        ('kompenzácia',   'O',             'U [kV]'),
        ('motoryVN',      'AD, AE',        '|I| [A], P_loss [kW]'),
        ('generatory',    'R, S, AC, AD',  'δ [°], Q_gen [Mvar], |I| [A], P_loss [kW]'),
        ('spinace',       'N',             '|I| [A]'),
    ],
    col_widths=[3.5, 3.5, 9.0])

add_heading(doc, '7.2 Samostatné výstupné hárky', level=2)
add_table(doc,
    ['Hárok', 'Obsah'],
    [
        ('Y_matica', 'Admitančná matica rozdelená na G = Re(Y) a B = Im(Y) so záhlavím uzlov.'),
        ('SC_matica','Skratová admitančná matica použitá pri Ik3″.'),
        ('napatia',  'Iteračný log napätí počas behu NR (iterácia, uzol, |V|, θ).'),
        ('epsilon',  'Iteračný log konvergencie (iterácia, max|ΔP|, max|ΔQ|, ε).'),
        ('report',   'Zoznam izolovaných prvkov siete.'),
        ('SLD',      'Jednolinková schéma s aktualizovanými hodnotami.'),
    ],
    col_widths=[3.5, 12.5])

add_heading(doc, '7.3 Aktualizácia jednolinkovej schémy (SLD)', level=2)
add_para(doc,
    'Modul modSLD.bas zapisuje výsledky priamo do jednolinkovej schémy podľa '
    'tzv. tagov. Tag je text v bunke schémy v tvare:')
p = doc.add_paragraph()
run = p.add_run('   <PREFIX>_<NÁZOV PRVKU>_<PREMENNÁ>_<SMER>')
run.font.name = 'Consolas'
run.font.size = Pt(10)
add_para(doc,
    'Príklad: tag N_Bratislava_V_R znamená „uzol Bratislava, premenná |V|, '
    'zapíš hodnotu napravo". Tag je v schéme bielou farbou textu (nevidno ho), '
    'po vyhodnotení sa do okolitej bunky zapíše naformátovaná hodnota s jednotkou.')

add_table(doc,
    ['Prefix', 'Typ prvku'],
    [
        ('N_',  'uzol'),
        ('V_',  'vedenie'),
        ('T_',  'transformátor'),
        ('R_',  'reaktor'),
        ('DR_', 'diferenciálny reaktor'),
        ('K_',  'kompenzácia'),
        ('M_',  'motor'),
        ('G_',  'generátor'),
        ('Q_',  'spínač'),
    ],
    col_widths=[2.5, 6.0])

add_para(doc,
    'Smer (R, L, U, D) určuje, ktorou stranou od tagu sa hodnota zapíše. '
    'Funkcie Odkryt_tagy / Skryt_tagy umožňujú dočasne odhaliť tagy zmenou '
    'farby textu na čiernu pre kontrolu zapojenia.')

add_placeholder(doc, 'hárok „SLD" s vyplnenými výsledkami v schéme')

doc.add_page_break()

# =======================================================
# 8. POSTUP VÝPOČTU
# =======================================================
add_heading(doc, '8. Postup výpočtu (procedúra runCALC)', level=1)
add_para(doc,
    'Hlavná procedúra runCALC v module modMain.bas tvorí jednotný vstupný bod '
    'volaný z tlačidla v hárku index. Beh je rozdelený do fáz, ktoré sa '
    'farebne indikujú v stavovom paneli (bunky I3:J8 hárka index).')

add_table(doc,
    ['Fáza', 'Bunka', 'Popis činnosti'],
    [
        ('1. Načítanie dát a topológia', 'I3/J3',
            'Načítanie všetkých vstupných hárkov do interných polí, '
            'BFS detekcia izolovaných prvkov.'),
        ('2. Tvorba matice',             'I4/J4',
            'BuildYBus (load-flow) alebo BuildShortCircuitMatrix (skraty).'),
        ('3a. Newton–Raphson',            'I5/J5/I6',
            'Iteračný cyklus, log napätí a ε. Iba ak mode = 1.'),
        ('3b. Skraty',                    'I7/J7',
            'Inverzia Y_sc a výpočet Ik3″. Iba ak mode = 2.'),
        ('4. Aktualizácia SLD',           'I8/J8',
            'Zápis výsledkov do jednolinkovej schémy podľa tagov.'),
    ],
    col_widths=[5.0, 2.0, 9.0])

add_para(doc, 'Stavový farebný kód (modul modProgress.bas):', bold=True)
add_bullet(doc, 'oranžová – fáza nebola spustená;')
add_bullet(doc, 'modrá – fáza prebieha;')
add_bullet(doc, 'zelená – fáza úspešne dokončená;')
add_bullet(doc, 'červená – chyba (zobrazený popis v stĺpci J);')
add_bullet(doc, 'sivá – fáza je v aktuálnom režime vypnutá (napr. NR pri skratovom režime).')

add_para(doc,
    'Pre dlhé výpočty je k dispozícii heartbeat (PhaseYield), ktorý raz za '
    'cca 200 ms aktualizuje časovač v bunke J5 a spracuje DoEvents (umožňuje '
    'reakciu na ESC a obnovenie UI). Stlačenie ESC spôsobí riadené prerušenie '
    's vrátením všetkých Excel nastavení (ScreenUpdating, výpočet, statusbar).')

add_placeholder(doc, 'stavový panel I3:J8 s farebnou indikáciou fáz')

# =======================================================
# 9. OBMEDZENIA A PRESNOSŤ
# =======================================================
add_heading(doc, '9. Obmedzenia, presnosť a požiadavky', level=1)

add_table(doc,
    ['Vlastnosť', 'Hodnota / popis'],
    [
        ('Platforma',           'Microsoft Excel 2010 a novší (Windows), VBA 7.x'),
        ('Reprezentácia čísel', 'IEEE 754 double (15–17 platných desatinných miest)'),
        ('Aritmetika',          'Natívne komplexné operácie (UDT Complex)'),
        ('Riešič sústavy',      'Gaussova eliminácia s čiastočným pivotovaním'),
        ('Max. veľkosť siete',  'Rádovo stovky uzlov (čas rastie ~n³ vďaka inverzii)'),
        ('Konvergencia NR',     'kvadratická; default ε = 10⁻⁶ p.u., max 20 iterácií'),
        ('Norma skratu',        'STN EN 60909 (c = 1.1 pre Ik3″ max)'),
        ('Jednotky výstupov',   'kV, A, MW, MVAr, kW (na vstupe Ω, S, MW, MVAr, kV)'),
    ],
    col_widths=[5.0, 11.0])

# =======================================================
# 10. PRÍLOHA – Mapovanie hárkov
# =======================================================
add_heading(doc, '10. Príloha – súhrnné mapovanie vstupov a výstupov', level=1)

add_para(doc,
    'Nasledujúca tabuľka súhrnne udáva, ktoré hárky program používa, či sú '
    'vstupné, výstupné, alebo zmiešané (vstup aj výstup):')
add_table(doc,
    ['Hárok',          'Smer',         'Krátky popis'],
    [
        ('index',         'Vstup',         'Riadiace parametre, výber režimu, stavový panel.'),
        ('data',          'Vstup',         'Globálne bázy (Sbase, napäťové hladiny VLevels).'),
        ('uzly',          'Vstup + Výstup','Definícia uzlov, výsledné |V|, θ, Ik3″.'),
        ('vedenia',       'Vstup + Výstup','π-model vedení, výsledné prúdy, toky, straty.'),
        ('transformatory','Vstup + Výstup','Trafá s prevodom, prúdy a toky na obidvoch stranách.'),
        ('reaktory',      'Vstup + Výstup','Sériové reaktory.'),
        ('dif_reaktory',  'Vstup + Výstup','Sekciovacie reaktory.'),
        ('kompenzácia',   'Vstup + Výstup','Priečne L/C kompenzácie.'),
        ('motoryVN',      'Vstup + Výstup','Asynchrónne motory VN.'),
        ('generatory',    'Vstup + Výstup','Synchrónne generátory (PQ alebo EMF).'),
        ('spinace',       'Vstup + Výstup','Spínacie prvky (vplyv na topológiu).'),
        ('Y_matica',      'Výstup',        'Admitančná matica G a B.'),
        ('SC_matica',     'Výstup',        'Skratová admitančná matica.'),
        ('napatia',       'Výstup',        'Iteračný log napätí.'),
        ('epsilon',       'Výstup',        'Iteračný log konvergencie.'),
        ('report',        'Výstup',        'Izolované prvky siete.'),
        ('SLD',           'Vstup + Výstup','Jednolinková schéma (tagy → hodnoty).'),
    ],
    col_widths=[3.5, 3.0, 9.5])

# záver
add_para(doc, '')
add_para(doc,
    'Tento dokument popisuje výpočtový model implementovaný v aktuálnej verzii '
    'modulov. Akýkoľvek zásah do zdrojového kódu (úprava vzorcov, pridanie '
    'nových typov prvkov, zmena štruktúry vstupných stĺpcov) je potrebné '
    'odzrkadliť aj v tejto dokumentácii.', italic=True, size=10)

# Save
out = '/home/user/excel/Dokumentacia_vypoctoveho_programu.docx'
doc.save(out)
print(f'Saved: {out}')
